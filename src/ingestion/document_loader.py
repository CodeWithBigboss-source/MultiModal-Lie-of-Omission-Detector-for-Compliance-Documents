"""
Document Ingestion Layer — Phase 2.

Accepts any document type the user uploads and produces a unified output:
    - extracted_text: str  → feeds into Step A (claim extraction)
    - extracted_images: dict[str, PIL.Image] → feeds into Step B (grounding)

Supported formats:
    - TXT  → direct read
    - DOCX → python-docx extracts paragraphs + embedded images
    - PDF  → pypdf for text-native pages, pytesseract OCR for scanned pages
    - JPG/PNG/WEBP → treated as pure image evidence, no text extraction
"""

import io
import uuid
from dataclasses import dataclass, field
from pathlib import Path

from PIL import Image


@dataclass
class DocumentContent:
    extracted_text: str = ""
    extracted_images: dict = field(default_factory=dict)  # image_id -> PIL.Image
    source_filename: str = ""
    page_count: int = 0
    extraction_notes: list = field(default_factory=list)


def _extract_from_txt(file_bytes: bytes) -> DocumentContent:
    text = file_bytes.decode("utf-8", errors="replace")
    return DocumentContent(
        extracted_text=text.strip(),
        page_count=1,
    )


def _extract_from_docx(file_bytes: bytes) -> DocumentContent:
    try:
        import docx
    except ImportError:
        return DocumentContent(
            extraction_notes=["python-docx not installed. Run: pip install python-docx"]
        )

    doc = docx.Document(io.BytesIO(file_bytes))
    text_parts = []
    images = {}

    # Extract paragraph text
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text.strip())

    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                text_parts.append(row_text)

    # Extract embedded images
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            try:
                img_bytes = rel.target_part.blob
                img = Image.open(io.BytesIO(img_bytes)).convert("RGB")
                img_id = f"docx_img_{uuid.uuid4().hex[:8]}"
                images[img_id] = img
            except Exception:
                pass

    return DocumentContent(
        extracted_text="\n".join(text_parts),
        extracted_images=images,
        page_count=1,
        extraction_notes=[f"Extracted {len(images)} embedded image(s) from DOCX."] if images else [],
    )


def _extract_from_pdf(file_bytes: bytes) -> DocumentContent:
    try:
        import pypdf
    except ImportError:
        return DocumentContent(
            extraction_notes=["pypdf not installed. Run: pip install pypdf"]
        )

    text_parts = []
    images = {}
    notes = []
    scanned_pages = []

    reader = pypdf.PdfReader(io.BytesIO(file_bytes))
    page_count = len(reader.pages)

    for page_num, page in enumerate(reader.pages):
        page_text = page.extract_text() or ""

        if len(page_text.strip()) > 50:
            # Text-native page
            text_parts.append(f"[Page {page_num + 1}]\n{page_text.strip()}")
        else:
            # Scanned page — needs OCR
            scanned_pages.append(page_num)

        # Extract embedded images from PDF page
        if hasattr(page, "images"):
            for img_obj in page.images:
                try:
                    img = Image.open(io.BytesIO(img_obj.data)).convert("RGB")
                    img_id = f"pdf_p{page_num + 1}_{uuid.uuid4().hex[:6]}"
                    images[img_id] = img
                except Exception:
                    pass

    # OCR scanned pages if any
    if scanned_pages:
        notes.append(
            f"{len(scanned_pages)} scanned page(s) detected. "
            "Attempting OCR — requires Tesseract installed on your system."
        )
        try:
            import pytesseract
            from pdf2image import convert_from_bytes

            pdf_images = convert_from_bytes(file_bytes, dpi=200)
            for page_num in scanned_pages:
                if page_num < len(pdf_images):
                    page_img = pdf_images[page_num]
                    ocr_text = pytesseract.image_to_string(page_img)
                    if ocr_text.strip():
                        text_parts.append(
                            f"[Page {page_num + 1} — OCR]\n{ocr_text.strip()}"
                        )
                    # Also keep the page image as visual evidence
                    img_id = f"pdf_scan_p{page_num + 1}"
                    images[img_id] = page_img.convert("RGB")
        except ImportError:
            notes.append(
                "OCR skipped — pdf2image or pytesseract not available. "
                "Install Tesseract from https://github.com/UB-Mannheim/tesseract/wiki "
                "then run: pip install pdf2image pytesseract"
            )
        except Exception as e:
            notes.append(f"OCR failed: {str(e)}")

    return DocumentContent(
        extracted_text="\n\n".join(text_parts),
        extracted_images=images,
        page_count=page_count,
        extraction_notes=notes,
    )


def _extract_from_image(file_bytes: bytes, filename: str) -> DocumentContent:
    """Pure image file — no text extraction, goes straight to visual evidence."""
    img = Image.open(io.BytesIO(file_bytes)).convert("RGB")
    img_id = f"upload_{Path(filename).stem}"
    return DocumentContent(
        extracted_text="",
        extracted_images={img_id: img},
        page_count=1,
        extraction_notes=["Image file — used as visual evidence only."],
    )


def load_document(file_bytes: bytes, filename: str) -> DocumentContent:
    """
    Main entry point. Call this with the raw bytes and filename of any uploaded file.
    Returns DocumentContent with extracted text and images ready for the pipeline.
    """
    ext = Path(filename).suffix.lower()
    content = None

    if ext == ".txt":
        content = _extract_from_txt(file_bytes)
    elif ext in (".docx", ".doc"):
        content = _extract_from_docx(file_bytes)
    elif ext == ".pdf":
        content = _extract_from_pdf(file_bytes)
    elif ext in (".jpg", ".jpeg", ".png", ".webp"):
        content = _extract_from_image(file_bytes, filename)
    else:
        content = DocumentContent(
            extraction_notes=[f"Unsupported file type: {ext}"]
        )

    content.source_filename = filename
    return content